from __future__ import annotations

import io
from collections.abc import Callable, Iterator

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell

Translator = Callable[[str, str, str], str]


def translate_docx(
    stream: io.BufferedReader | io.BytesIO, *, source: str, target: str, translator: Translator
) -> bytes:
    document: DocumentType = Document(stream)
    for paragraph in _iter_paragraphs(document):
        _translate_paragraph(paragraph, source=source, target=target, translator=translator)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _iter_paragraphs(document: DocumentType) -> Iterator:
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            yield from header.paragraphs
            for table in header.tables:
                yield from _iter_table_paragraphs(table)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            yield from footer.paragraphs
            for table in footer.tables:
                yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table) -> Iterator:
    for row in table.rows:
        for cell in row.cells:
            cell: _Cell
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _translate_paragraph(paragraph, *, source: str, target: str, translator: Translator) -> None:
    text = paragraph.text
    if not text.strip():
        return

    translated = translator(text, source, target)
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(translated)
        return

    runs[0].text = translated
    for run in runs[1:]:
        run.text = ""
