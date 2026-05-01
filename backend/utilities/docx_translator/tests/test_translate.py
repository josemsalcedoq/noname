import io
import json
from pathlib import Path

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client
from docx import Document
from pytest_bdd import given, parsers, scenarios, then, when

scenarios(str(Path(__file__).parent / "features" / "translate.feature"))


@pytest.fixture
def state():
    return {"upload": None, "source": "en", "target": "es", "response": None}


@pytest.fixture
def client():
    return Client()


def _make_docx(paragraphs: list[str]) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _read_paragraphs(content: bytes) -> list[str]:
    document = Document(io.BytesIO(content))
    return [p.text for p in document.paragraphs]


@given(parsers.parse("a docx with paragraphs {paragraphs}"))
def _given_docx(state, paragraphs):
    parsed = json.loads(paragraphs)
    state["upload"] = SimpleUploadedFile(
        "input.docx",
        _make_docx(parsed),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@given("an upload that is not a docx file")
def _given_non_docx(state):
    state["upload"] = SimpleUploadedFile("input.txt", b"plain text", content_type="text/plain")


@when(parsers.parse('the client uploads it with source "{source}" and target "{target}"'))
def _upload(state, client, fake_translate, source, target):
    state["source"] = source
    state["target"] = target
    state["response"] = client.post(
        "/api/docx-translator/translate",
        data={"file": state["upload"], "source": source, "target": target},
        format="multipart",
    )


@when("the client posts to docx translate with no file")
def _upload_no_file(state, client):
    state["response"] = client.post(
        "/api/docx-translator/translate",
        data={"source": "en", "target": "es"},
        format="multipart",
    )


@then(parsers.parse("the response status is {code:d}"))
def _status(state, code):
    assert state["response"].status_code == code


@then(parsers.parse("the response is a docx with paragraphs {paragraphs}"))
def _docx_paragraphs(state, paragraphs):
    expected = json.loads(paragraphs)
    actual = _read_paragraphs(b"".join(state["response"].streaming_content))
    assert actual == expected


@then(parsers.parse('the response filename ends with "{suffix}"'))
def _filename(state, suffix):
    disposition = state["response"]["Content-Disposition"]
    assert suffix in disposition


@then(parsers.parse('the response error code is "{code}"'))
def _error_code(state, code):
    assert state["response"].json()["error"] == code
